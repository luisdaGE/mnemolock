from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("docs/MnemoLock_Brief_Claude_Codex.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 35, 52)
MUTED = RGBColor(91, 103, 115)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BORDER = "C9D3DF"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tbl_pr.append(ind)
    ind.set(qn("w:w"), str(indent_dxa))
    ind.set(qn("w:type"), "dxa")


def style_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "MnemoLock | Brief para mejora con Claude y ejecucion con Codex"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Documento generado desde el repositorio local"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("MnemoLock")
    set_run_font(run, size=26, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Brief completo de flujo, negocio, arquitectura y mejoras")
    set_run_font(run, size=14, color=MUTED)

    rows = [
        ("Objetivo", "Pasar a Claude una vision completa de la app y recibir mejoras concretas."),
        ("Uso esperado", "Claude propone mejoras; Codex implementa cambios verificables en este repositorio."),
        ("Estado", "MVP web/PWA con bloqueo simulado, quiz de desbloqueo y Supabase preparado."),
        ("Fecha de referencia", "9 de junio de 2026"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.45)
        cells[1].width = Inches(4.95)
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)
        set_cell_shading(cells[0], LIGHT_GRAY)
        cells[0].paragraphs[0].runs[0].bold = True


def add_para(doc, text, style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        if widths:
            cell.width = Inches(widths[idx])
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        set_cell_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = DARK_BLUE

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            cell.text = text
            if widths:
                cell.width = Inches(widths[idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = INK
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FB")
    set_cell_border(cell, "D3DCE8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.add_run(body)
    doc.add_paragraph()


def build():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    add_callout(
        doc,
        "Idea central",
        "MnemoLock es una PWA de estudio que bloquea la sesion hasta que el usuario demuestra aprendizaje con preguntas del tema. El desbloqueo no depende de fuerza de voluntad, sino de memoria activa.",
    )
    add_bullets(
        doc,
        [
            "Promesa: convertir cada intento de salir de una distraccion en una microevaluacion util.",
            "MVP actual: experiencia web responsive, rutas de producto, demo de bloqueo, quiz configurable, modo estricto, login Supabase preparado y ruta de fuentes.",
            "Diferenciador: preguntas enlazadas a fragmentos citables del material del estudiante para evitar IA opaca o bancos genericos.",
            "Limite tecnico: la PWA no puede bloquear otras apps a nivel sistema operativo; el bloqueo real requiere app nativa posterior.",
        ]
    )

    doc.add_heading("2. Flujo principal del usuario", level=1)
    add_table(
        doc,
        ["Paso", "Pantalla / modulo", "Que ocurre", "Dato clave"],
        [
            ("1", "Estudio / Dashboard", "El usuario elige materia, duracion y modo estricto.", "study_set, unlockTarget, cooldownMinutes"),
            ("2", "Sesion bloqueada", "La app entra en fullscreen, activa wake lock si el navegador lo permite y avisa antes de cerrar.", "SessionState = locked"),
            ("3", "Quiz", "Al terminar el tiempo, el usuario responde preguntas del tema.", "answers, activeQuestion, correctCount"),
            ("4", "Desbloqueo", "Si alcanza los aciertos requeridos, sale de fullscreen y la sesion queda desbloqueada.", "SessionState = unlocked"),
            ("5", "Penalizacion", "Si falla, modo estricto reinicia bloqueo con cooldown; sin estricto se queda en quiz.", "cooldownMinutes x 60"),
        ],
        widths=[0.55, 1.55, 3.1, 1.2],
    )
    add_para(
        doc,
        "Flujo visual resumido: Elegir tema -> Iniciar bloqueo -> Esperar foco -> Responder quiz -> Desbloquear si aprende -> Penalizar si falla.",
    )

    doc.add_heading("3. Rutas y experiencia actual", level=1)
    add_table(
        doc,
        ["Ruta", "Funcion", "Estado actual"],
        [
            ("/", "Home de estudio y demo", "Explica la promesa, muestra vista rapida e inicia demo."),
            ("/login", "Acceso", "Email/password y Google OAuth conectados a Supabase cuando existen variables de entorno."),
            ("/sources", "Fuentes", "Selector local de PDF/TXT/MD y pipeline visual de extraccion, filtrado, preguntas y citas."),
            ("/dashboard", "Espacio de estudio", "Guia paso a paso, controles de sesion y CTA para subir apuntes."),
            ("/strategy", "Producto", "Diferenciadores, debilidades del mercado y modelo de negocio."),
            ("/settings", "Ajustes", "Estructura futura para bloqueo, notificaciones, datos y permisos nativos."),
            ("/info/:slug", "Contenido informativo", "Paginas educativas sobre metodo, estudiantes y escuelas."),
        ],
        widths=[1.1, 1.8, 3.5],
    )

    doc.add_heading("4. Modelo mental del producto", level=1)
    add_bullets(
        doc,
        [
            "El bloqueo no debe sentirse como castigo permanente: es una puerta cognitiva temporal.",
            "El estudiante recupera acceso cuando demuestra dominio, no cuando negocia con la app.",
            "La IA debe ser verificable: cada pregunta generada tiene fuente, fragmento y pagina cuando aplique.",
            "El MVP debe preservar una experiencia simple; las funciones avanzadas no deben esconder el boton principal: estudiar.",
            "El camino nativo debe tratarse como fase posterior, no como promesa falsa de la PWA.",
        ]
    )

    doc.add_heading("5. Arquitectura tecnica actual", level=1)
    add_table(
        doc,
        ["Capa", "Implementacion", "Observaciones"],
        [
            ("Frontend", "React 18 + TypeScript + Vite", "App responsive con CSS propio y lucide-react."),
            ("Rutas", "react-router-dom", "Shell global en src/app y pantallas por feature."),
            ("Estado de sesion", "useStudySession", "Controla selectedSet, timer, quiz, fullscreen, wake lock y penalizacion."),
            ("Datos demo", "src/data/studySets.ts", "Materias iniciales: Biologia celular, Historia de Mexico y Algebra esencial."),
            ("Auth", "Supabase client", "Email/password y Google OAuth; fallback si faltan env vars."),
            ("Base de datos", "supabase/schema.sql", "RLS, perfiles, ajustes, fuentes, chunks, sesiones e intentos."),
            ("Deploy", "Vercel", "Scripts listos: dev, typecheck, build, preview."),
        ],
        widths=[1.35, 2.15, 2.9],
    )

    doc.add_heading("6. Modelo de datos propuesto", level=1)
    add_table(
        doc,
        ["Tabla", "Proposito"],
        [
            ("profiles", "Perfil y rol: student, guardian, teacher o admin."),
            ("user_settings", "Aciertos por defecto, cooldown, modo estricto y preferencias."),
            ("study_sets", "Materias/conjuntos de estudio por usuario."),
            ("questions", "Preguntas, opciones, respuesta, dificultad, explicacion y source_chunk_id."),
            ("study_sources", "Archivos cargados por usuario: nombre, ruta, MIME y estado."),
            ("source_chunks", "Fragmentos citables con indice, contenido, pagina y metadata."),
            ("focus_sessions", "Ciclo de bloqueo: duracion, estado, inicio, fin y metadata."),
            ("quiz_attempts", "Resultado agregado de cada intento de desbloqueo."),
            ("question_attempts", "Respuesta seleccionada y acierto/error por pregunta."),
        ],
        widths=[1.65, 4.75],
    )
    add_callout(
        doc,
        "Decision critica de confianza",
        "No generar preguntas sueltas sin source_chunk_id cuando vengan de materiales del usuario. La trazabilidad es parte del producto, no solo una mejora tecnica.",
    )

    doc.add_heading("7. Modelo de negocio", level=1)
    add_table(
        doc,
        ["Linea", "Oferta", "Razon estrategica"],
        [
            ("Gratis", "Sesiones, bancos manuales, rachas y estadisticas basicas.", "Reduce friccion y valida habito."),
            ("Pro individual", "IA desde apuntes, fotos, PDFs y videos; repasos espaciados; analiticas.", "Monetiza el valor principal: convertir material propio en desbloqueo confiable."),
            ("Escuelas", "Panel docente, retos por grupo, reportes agregados, privacidad y SSO.", "Expande a B2B con valor institucional."),
            ("Marketplace", "Packs verificados por docentes y creadores.", "Crea inventario educativo y revenue share."),
            ("B2B2C", "Convenios con academias, preparatorias y plataformas de cursos.", "Canal de adquisicion y retencion por cohortes."),
        ],
        widths=[1.15, 2.65, 2.6],
    )

    doc.add_heading("8. Oportunidad de mercado y diferenciadores", level=1)
    add_bullets(
        doc,
        [
            "Bloqueadores tradicionales dependen de autocontrol y pueden ser faciles de rodear.",
            "Apps de estudio suelen usar bancos genericos que no coinciden con el examen real.",
            "Herramientas de IA educativas fallan cuando no muestran fuente ni explican de donde sale la pregunta.",
            "MnemoLock combina foco, recuperacion activa y trazabilidad: tres problemas en un flujo corto.",
        ]
    )

    doc.add_heading("9. Roadmap recomendado", level=1)
    add_numbers(
        doc,
        [
            "Persistir sesiones reales en Supabase y separar datos demo de datos del usuario.",
            "Crear CRUD de materias y preguntas con validaciones y estados vacios profesionales.",
            "Subir PDFs a Supabase Storage, insertar study_sources y procesar texto server-side.",
            "Partir texto en source_chunks y generar preguntas con source_chunk_id obligatorio.",
            "Guardar quiz_attempts y question_attempts para analiticas y repaso espaciado.",
            "Agregar dashboard de progreso por materia, debilidad, racha y confiabilidad de fuentes.",
            "Convertir PWA en app nativa con Capacitor y modulos iOS/Android para bloqueo real.",
            "Construir modo escuela: grupos, docentes, assignments, reportes agregados y SSO.",
        ]
    )

    doc.add_heading("10. Riesgos y limites", level=1)
    add_table(
        doc,
        ["Riesgo", "Impacto", "Mitigacion"],
        [
            ("Bloqueo web limitado", "La PWA no puede bloquear otras apps del celular.", "Comunicarlo con honestidad y planear fase nativa."),
            ("IA sin fuente", "Preguntas incorrectas reducen confianza.", "Exigir citas y revisar fragmentos antes de aceptar."),
            ("Friccion al estudiar", "El usuario abandona si configurar toma demasiado.", "Primer flujo con demo, defaults y subida simple."),
            ("Privacidad escolar", "Datos de menores o grupos requieren cuidado.", "RLS, politicas claras, roles y reportes agregados."),
            ("Complejidad de permisos nativos", "iOS/Android requieren caminos distintos.", "Aislar logica de sesion y planear adaptadores nativos."),
        ],
        widths=[1.75, 2.25, 2.4],
    )

    doc.add_heading("11. Prompt listo para Claude", level=1)
    add_para(
        doc,
        "Copia y pega este bloque en Claude para pedirle una mejora estructurada del producto:",
    )
    prompt = (
        "Actua como estratega de producto, arquitecto tecnico y UX lead. Estoy construyendo MnemoLock, "
        "una PWA de estudio que bloquea una sesion hasta que el usuario demuestra aprendizaje con un quiz. "
        "El MVP actual tiene React + TypeScript + Vite, Supabase preparado, rutas de estudio/fuentes/dashboard/"
        "producto/ajustes, bloqueo simulado con fullscreen/wake lock/beforeunload, quiz configurable y modo estricto. "
        "La promesa es desbloquear por memoria activa, no por fuerza de voluntad. La futura ventaja es subir PDFs/"
        "apuntes, extraer fragmentos citables y generar preguntas enlazadas a source_chunk_id. Revisa el flujo, "
        "modelo de negocio, UX, datos, riesgos, monetizacion y roadmap. Devuelveme: 1) mejoras priorizadas por impacto, "
        "2) cambios concretos en pantallas, 3) cambios de arquitectura/base de datos, 4) experimentos de negocio, "
        "5) texto de producto mejorado, 6) lista exacta de tareas para que Codex implemente en el repo. No propongas "
        "bloqueo real web de otras apps; si hablas de bloqueo real, marcala como fase nativa con iOS/Android."
    )
    add_callout(doc, "Prompt para Claude", prompt)

    doc.add_heading("12. Formato de respuesta que Claude debe devolver", level=1)
    add_bullets(
        doc,
        [
            "Prioridad: P0, P1 o P2.",
            "Tipo: producto, UX, frontend, datos, backend, negocio o copy.",
            "Archivos probables a tocar.",
            "Cambio esperado en una frase.",
            "Criterios de aceptacion verificables.",
            "Riesgos o dependencias.",
        ]
    )

    doc.add_heading("13. Instrucciones para Codex al implementar", level=1)
    add_bullets(
        doc,
        [
            "Leer primero README.md, PRODUCT_STRATEGY.md, docs/PROJECT_STRUCTURE.md, docs/SCALABILITY_PLAN.md y supabase/schema.sql.",
            "Mantener la arquitectura por features; no volver a concentrar logica en App.tsx.",
            "No prometer bloqueo real del sistema desde la PWA; cualquier texto debe distinguir demo web vs fase nativa.",
            "Preservar la trazabilidad: preguntas generadas desde fuentes deben enlazarse a source_chunks.",
            "Antes de cambios grandes, correr npm run typecheck y npm run build si las dependencias estan instaladas.",
            "Agregar tests en utils/study.ts y luego cubrir LockPanel / flujo E2E cuando el alcance toque sesion.",
            "Mantener UI simple, responsive y enfocada en iniciar una sesion rapido.",
        ]
    )

    doc.add_heading("14. Tareas iniciales sugeridas para Codex", level=1)
    add_table(
        doc,
        ["Prioridad", "Tarea", "Archivos probables", "Aceptacion"],
        [
            ("P0", "Persistir focus_sessions y quiz_attempts cuando Supabase este configurado.", "useStudySession, lib/supabase, types/database", "Una sesion crea/actualiza registros y maneja fallback demo."),
            ("P0", "Crear CRUD basico de study_sets/questions.", "features/dashboard, data, supabase helpers", "Usuario autenticado puede crear materia y preguntas propias."),
            ("P1", "Implementar subida real a Supabase Storage.", "SourcesPage, lib/supabase, schema/storage docs", "Archivo sube, crea study_sources y muestra estado."),
            ("P1", "Crear Edge Function o worker para chunking.", "supabase/functions, sourcePipeline", "source_chunks quedan asociados con pagina/metadata."),
            ("P1", "Mejorar dashboard con progreso real.", "DashboardPage, query layer", "Muestra sesiones, intentos, aciertos y materias debiles."),
            ("P2", "Preparar arquitectura nativa Capacitor.", "docs, config, future native adapters", "Documento tecnico separa PWA y bloqueo nativo."),
        ],
        widths=[0.75, 2.25, 1.75, 1.65],
    )

    doc.add_heading("15. Glosario operativo", level=1)
    add_table(
        doc,
        ["Termino", "Definicion"],
        [
            ("Memoria activa", "Recuperar informacion sin verla; base pedagogica del desbloqueo."),
            ("Modo estricto", "Si el usuario falla el quiz, vuelve a bloqueo con cooldown."),
            ("source_chunk", "Fragmento de material usado como evidencia para una pregunta."),
            ("unlockTarget", "Cantidad minima de aciertos para desbloquear."),
            ("focus_session", "Registro persistente de una sesion de enfoque/bloqueo."),
            ("quiz_attempt", "Intento de desbloqueo con resultado passed/failed."),
        ],
        widths=[1.7, 4.7],
    )

    doc.add_heading("16. Cierre", level=1)
    add_para(
        doc,
        "La mejora mas importante no es agregar muchas funciones: es cerrar el ciclo entre material propio, pregunta verificable, intento de desbloqueo y aprendizaje medible. Todo cambio deberia reforzar ese ciclo.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
